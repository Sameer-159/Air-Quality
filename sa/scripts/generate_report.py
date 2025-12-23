from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a professional 11-page report with placeholders
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
import os


def load_dataset(path='data/AirQualityUCI.csv'):
    # Read dataset using same preprocessing as fuzzy_system
    if not os.path.exists(path):
        return None
    try:
        df = pd.read_csv(path, sep=';', decimal=',')
        df.columns = df.columns.str.strip()
        df = df.replace(-200, pd.NA)
        # convert numeric columns
        for col in df.columns:
            if col not in ['Date', 'Time']:
                df[col] = pd.to_numeric(df[col], errors='coerce')
        return df
    except Exception:
        return None


def stats_for_columns(df, cols):
    rows = []
    for c in cols:
        if c in df.columns:
            series = df[c].dropna()
            rows.append((c, int(len(series)), float(series.min()), float(series.max()), float(series.mean()), float(series.std())))
        else:
            rows.append((c, 0, None, None, None, None))
    return rows


def make_report(output_path='report.docx'):
    doc = Document()

    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Fuzzy Air Quality Assessment System')
    run.bold = True
    run.font.size = Pt(28)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Comprehensive Analysis and Membership Function Visualizations')
    run.font.size = Pt(14)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Author: ____________________')
    run.font.size = Pt(12)

    doc.add_page_break()

    # Page 2: Executive Summary (detailed but concise)
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'This report presents a fuzzy-logic based system for assessing urban air quality using the Air Quality UCI dataset. '
        'We developed membership functions for key pollutants (CO, NOx/NO2, O3), environmental parameters (temperature, humidity), '
        'and derived an AQI output. The fuzzy approach captures uncertainty and partial memberships to provide a more interpretable '
        'estimate compared to crisp baselines.'
    )
    doc.add_paragraph(
        'Key findings (placeholders):\n'
        '- The fuzzy AQI aligns with EPA-like crisp calculations on average, while providing smoother category transitions.\n'
        '- Membership functions highlight sensitive ranges where small changes in pollutant concentration produce meaningful AQI shifts.\n'
        '- Advanced comparisons (MAE, RMSE, categorical accuracy) should be inserted in the Performance section.'
    )
    doc.add_page_break()

    # Page 3: Dataset Description (with stats table)
    df = load_dataset()
    doc.add_heading('Dataset Description', level=1)
    doc.add_paragraph('Dataset used: AirQualityUCI.csv (Hourly sensor readings from a multisensor device)')
    doc.add_paragraph('The dataset contains multi-sensor measurements along with meteorological parameters. The following table summarizes available data ranges and sample counts:')

    cols = ['CO(GT)', 'NO2(GT)', 'PT08.S5(O3)', 'T', 'RH', 'AH']
    stats = stats_for_columns(df, cols) if df is not None else []

    table = doc.add_table(rows=1, cols=6)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Parameter'
    hdr_cells[1].text = 'Samples'
    hdr_cells[2].text = 'Min'
    hdr_cells[3].text = 'Max'
    hdr_cells[4].text = 'Mean'
    hdr_cells[5].text = 'Std'

    for row in stats:
        cells = table.add_row().cells
        cells[0].text = str(row[0])
        cells[1].text = str(row[1])
        cells[2].text = '-' if row[2] is None else f"{row[2]:.3f}"
        cells[3].text = '-' if row[3] is None else f"{row[3]:.3f}"
        cells[4].text = '-' if row[4] is None else f"{row[4]:.3f}"
        cells[5].text = '-' if row[5] is None else f"{row[5]:.3f}"

    doc.add_paragraph('\nNote: missing values were filtered. For noisy sensor fields the raw sensor voltages are also available in the dataset.')
    doc.add_page_break()

    # Page 4: Methodology (extended)
    doc.add_heading('Methodology', level=1)
    doc.add_paragraph(
        'System design: The fuzzy system uses antecedent variables for pollutant concentrations and environmental conditions, and a consequent variable for AQI. '
        'Each input is described by linguistically meaningful terms (for example CO: Very Low, Low, Moderate, High, Very High). Membership functions were defined using triangular and trapezoidal functions to reflect sensor and health thresholds.'
    )
    doc.add_paragraph(
        'Rule base: A set of human-interpretable rules combine pollutant levels and contextual features (temperature, humidity) to infer AQI categories. '
        'Rules include single-pollutant dominance (e.g., very high CO leads to Hazardous AQI) and combined pollutant interactions (e.g., moderate NO2 with high O3 increases the AQI). '
        'The system uses Mamdani-style inference with centroid defuzzification.'
    )
    doc.add_paragraph('Placeholder: system architecture diagram and pseudocode for rule evaluation.')
    doc.add_page_break()

    # Page 5: Membership Functions Overview (detailed guidance)
    doc.add_heading('Membership Functions - Overview', level=1)
    doc.add_paragraph(
        'This section presents the membership functions for primary pollutants. For each parameter, include a plot showing the universe of discourse and overlaid membership curves. '
        'Recommended plots to include: CO (GT), NO2 (GT), O3 (Sensor), and AQI output membership functions.'
    )
    for name in ['CO (Ground Truth)', 'NO2 (Ground Truth)', 'O3 (Sensor)', 'AQI (Output)']:
        doc.add_heading(name, level=2)
        doc.add_paragraph('[Insert plot image here: {}]'.format(name)).italic = True
        doc.add_paragraph('Caption: Describe the key ranges and how they map to linguistic terms. For example: "CO > 10 ppm maps strongly to Very High, which increases AQI rapidly."')
    doc.add_page_break()

    # Page 6: Additional Membership Functions
    doc.add_heading('Membership Functions - Additional Parameters', level=1)
    doc.add_paragraph('Include membership functions for sensor-specific readings and supporting variables: CO sensor, NMHC, Benzene, NOx, Temperature, Humidity.')
    for name in ['CO Sensor', 'NMHC (GT)', 'Benzene (GT)', 'NOx (GT)', 'Temperature', 'Humidity']:
        doc.add_heading(name, level=2)
        doc.add_paragraph('[Insert plot image here: {}]'.format(name)).italic = True
    doc.add_page_break()

    # Page 7: Advanced Visualizations
    doc.add_heading('Advanced Visualizations', level=1)
    doc.add_paragraph('Advanced visualizations help diagnose system behaviour and interpretability:')
    doc.add_paragraph('- Radar charts comparing fuzzy vs crisp performance metrics across multiple dimensions.')
    doc.add_paragraph('- Bar charts showing rule activations for representative cases to explain decisions.')
    doc.add_paragraph('[Insert radar chart image]').italic = True
    doc.add_paragraph('[Insert rule activations bar chart]').italic = True
    doc.add_page_break()

    # Page 8: Results - Sample Assessments
    doc.add_heading('Results - Sample Assessments', level=1)
    doc.add_paragraph('Provide several worked examples: table of input values (CO, NO2, O3, Temp, Humidity), fuzzy AQI, crisp AQI, and membership values for key terms. Use plots to show how membership degrees lead to the final decision.')
    doc.add_paragraph('[Insert sample assessment plots and tables]').italic = True
    doc.add_page_break()

    # Page 9: Performance Comparison and Metrics
    doc.add_heading('Performance Comparison and Metrics', level=1)
    doc.add_paragraph('Quantitative evaluation should include: MAE, RMSE for AQI values, categorical accuracy (AQI bands), F1-score, and satisfaction (percentage within acceptable error). Include confidence intervals and discuss limits of evaluation given sensor noise and missing data.')
    doc.add_paragraph('[Insert performance metrics table and comparison charts]').italic = True
    doc.add_page_break()

    # Page 10: Discussion
    doc.add_heading('Discussion', level=1)
    doc.add_paragraph('Interpretation: Discuss where fuzzy logic adds value — handling ambiguous ranges, providing interpretable rules, and smoothing transitions between categories. Address limitations: dataset biases, sensor calibration differences, and seasonal variations.')
    doc.add_paragraph('Recommendations: Suggest further validation with co-located reference instruments, refinement of membership thresholds using domain experts, and deployment considerations (real-time processing, drifting sensors).')
    doc.add_page_break()

    # Page 11: Conclusions and Appendices
    doc.add_heading('Conclusions and Next Steps', level=1)
    doc.add_paragraph('Conclusions: Summarize the main takeaways and potential impact for air quality monitoring and public health communications.')
    doc.add_paragraph('Next steps: list implementation, validation, and monitoring steps; include potential integration with IoT dashboards and alerting systems.')
    doc.add_paragraph('Appendix A: Code references - see repository files: app.py, fuzzy_system.py, static/js/app.js')
    doc.add_paragraph('Appendix B: API endpoints - /assess, /enhanced_assess, /membership_functions, /compare, /advanced_compare, /test_assessment')

    # Save
    doc.save(output_path)
    print('Report generated:', output_path)


if __name__ == '__main__':
    make_report()
